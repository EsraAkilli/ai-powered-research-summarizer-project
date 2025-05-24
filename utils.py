import io
import os
import tempfile
from docx import Document
from docx.shared import Inches
import markdown
from bs4 import BeautifulSoup
import streamlit as st


def convert_md_to_docx_memory(md_content: str) -> bytes:
    """
    Markdown içeriğini memory'de DOCX bytes olarak dönüştürür.
    Geçici dosya oluşturmadan çalışır.
    """
    try:
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')

        doc = Document()

        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                heading = doc.add_heading(element.get_text().strip(), level=level)
            elif element.name == 'p':
                doc.add_paragraph(element.get_text().strip())
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    doc.add_paragraph(li.get_text().strip(), style='List Bullet')

        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        return doc_bytes.getvalue()

    except Exception as e:
        st.error(f"DOCX dönüştürme hatası: {str(e)}")
        return None


def convert_md_to_docx_temp(md_content: str) -> bytes:
    """
    Markdown içeriğini geçici dosya kullanarak DOCX bytes olarak dönüştürür.
    Fallback method olarak kullanılır.
    """
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
            temp_md.write(md_content)
            temp_md_path = temp_md.name

        with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as temp_docx:
            temp_docx_path = temp_docx.name

        doc = Document()

        lines = md_content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], 4)
            elif line and not line.startswith('#'):
                doc.add_paragraph(line)

        doc.save(temp_docx_path)

        with open(temp_docx_path, 'rb') as f:
            docx_bytes = f.read()

        os.unlink(temp_md_path)
        os.unlink(temp_docx_path)

        return docx_bytes

    except Exception as e:
        st.error(f"DOCX dönüştürme hatası: {str(e)}")
        return None


def convert_md_to_docx(md_content: str) -> bytes:
    """
    Ana dönüştürme fonksiyonu. Önce memory-based method'u dener,
    başarısız olursa temp file method'unu kullanır.
    """
    result = convert_md_to_docx_memory(md_content)
    if result is not None:
        return result

    return convert_md_to_docx_temp(md_content)


def safe_filename(query: str, max_length: int = 50) -> str:
    """
    Güvenli dosya adı oluşturur
    """
    safe_chars = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789-_"
    filename = ''.join(c if c in safe_chars else '_' for c in query)

    if len(filename) > max_length:
        filename = filename[:max_length]

    if not filename.strip('_'):
        filename = "research_summary"

    return filename